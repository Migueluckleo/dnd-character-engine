"""
build_behavioral_design.py
Genera behavioral design.docx a partir de docs/behavioral_design.md.
La fuente de verdad es el .md; el .docx es el artefacto de salida.

Uso:
    python3 scripts/build_behavioral_design.py
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Rutas ──────────────────────────────────────────────────────────────────
SCRIPT_DIR  = Path(__file__).parent
PROJECT_DIR = SCRIPT_DIR.parent
MD_PATH     = PROJECT_DIR / "docs" / "behavioral_design.md"
DOCX_OUT    = PROJECT_DIR / "behavioral design.docx"

# ── Colores temáticos ──────────────────────────────────────────────────────
C_DARK   = RGBColor(0x2C, 0x1A, 0x0E)
C_ACCENT = RGBColor(0x8B, 0x45, 0x13)
C_MUTED  = RGBColor(0x6B, 0x5A, 0x4E)
C_BODY   = RGBColor(0x1A, 0x1A, 0x1A)
C_GREEN  = RGBColor(0x22, 0x86, 0x34)
C_AMBER  = RGBColor(0xB4, 0x5A, 0x00)
C_RED    = RGBColor(0xC0, 0x20, 0x20)

STATUS_COLORS = {"✅": C_GREEN, "⚠️": C_AMBER, "❌": C_RED}


# ── Helpers ────────────────────────────────────────────────────────────────
def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Georgia"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    sizes  = {1: 18, 2: 15, 3: 13, 4: 11}
    colors = {1: C_DARK, 2: C_ACCENT, 3: C_ACCENT, 4: C_MUTED}
    set_font(run, size=sizes.get(level, 11), bold=True, color=colors.get(level, C_BODY))
    return p


def add_body(doc, text, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, italic=italic, color=color or C_BODY)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        set_font(run, size=10, bold=(i % 2 == 1), color=C_BODY)
    return p


def add_cover(doc, md_text):
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(40)
    r = title_p.add_run("BEHAVIORAL DESIGN")
    set_font(r, size=28, bold=True, color=C_DARK)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_p.add_run("DnD Character Engine")
    set_font(r2, size=16, color=C_ACCENT)

    tag_p = doc.add_paragraph()
    tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = tag_p.add_run("Jobs to Be Done · Historias de usuario · Criterios de aceptación")
    set_font(r3, size=11, italic=True, color=C_MUTED)

    version = "1.0"
    date    = "Mayo 2026"
    vm = re.search(r'\|\s*Versión\s*\|\s*(.+?)\s*\|', md_text)
    dm = re.search(r'\|\s*Última actualización\s*\|\s*(.+?)\s*\|', md_text)
    if vm: version = vm.group(1).strip()
    if dm: date    = dm.group(1).strip()

    ver_p = doc.add_paragraph()
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver_p.paragraph_format.space_before = Pt(8)
    r4 = ver_p.add_run(f"Versión {version}  ·  {date}")
    set_font(r4, size=10, color=C_MUTED)

    doc.add_page_break()


def add_jtbd_callout(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f"JOB TO BE DONE   {text}")
    set_font(run, size=10, italic=True, color=C_ACCENT)


def add_us_header(doc, us_code, title, story):
    add_heading(doc, f"{us_code} — {title}", level=3)
    if story:
        add_body(doc, story, italic=True, color=C_MUTED)


def add_status_line(doc, status_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f"Estado: {status_text}")
    color = C_MUTED
    for emoji, c in STATUS_COLORS.items():
        if emoji in status_text:
            color = c
            break
    set_font(run, size=10, bold=True, color=color)


# ── Parser principal ────────────────────────────────────────────────────────
def parse_and_render(doc, md_text):
    lines = md_text.splitlines()
    i = 0
    in_code_block = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code blocks
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue
        if in_code_block:
            i += 1
            continue

        # Blockquote / callout lines
        if stripped.startswith("> ") or stripped == ">":
            text = stripped.lstrip("> ").strip()
            if "JOB TO BE DONE" in text:
                text = re.sub(r'\*?\*?JOB TO BE DONE\*?\*?\s*[:\—\-]?\s*', '', text).strip()
                add_jtbd_callout(doc, text)
            elif text:
                add_body(doc, text, italic=True, color=C_MUTED)
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^---+$', stripped) or re.match(r'^\*\*\*+$', stripped):
            i += 1
            continue

        # H1 — skip (it's the doc title, already on cover)
        if re.match(r'^# [^#]', stripped):
            i += 1
            continue

        # H2
        if stripped.startswith("## "):
            text = stripped[3:].strip()
            # Skip version/historial tables
            if any(k in text for k in ("Versión", "Historial", "historial")):
                i += 1
                continue
            add_heading(doc, text, level=2)
            i += 1
            continue

        # H3
        if stripped.startswith("### "):
            text = stripped[4:].strip()
            us_match = re.match(r'^(US-[\w/\s\-]+?)\s*—\s*(.+)$', text)
            if us_match:
                us_code = us_match.group(1).strip()
                title   = us_match.group(2).strip()
                story = ""
                if i + 1 < len(lines):
                    next_l = lines[i + 1].strip()
                    if next_l.startswith("*") and next_l.endswith("*") and not next_l.startswith("**"):
                        story = next_l.strip("*").strip()
                        i += 1
                add_us_header(doc, us_code, title, story)
            else:
                add_heading(doc, text, level=3)
            i += 1
            continue

        # H4
        if stripped.startswith("#### "):
            add_heading(doc, stripped[5:].strip(), level=4)
            i += 1
            continue

        # Status line
        if stripped.startswith("**Estado:**"):
            status_text = stripped.replace("**Estado:**", "").strip()
            add_status_line(doc, status_text)
            i += 1
            continue

        # AC / criterios header
        if stripped == "**Criterios de aceptación:**":
            p = doc.add_paragraph()
            r = p.add_run("Criterios de aceptación:")
            set_font(r, size=10, bold=True, color=C_ACCENT)
            i += 1
            continue

        # Bullet points
        if re.match(r'^[-*]\s', stripped):
            text = stripped[2:].strip()
            if re.match(r'^[\|\-\s:]+$', text):
                i += 1
                continue
            add_bullet(doc, text)
            i += 1
            continue

        # Table rows
        if stripped.startswith("|"):
            if re.match(r'^[\|\-\s:]+$', stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            header_cells = {"Campo", "Contexto", "Capa", "Fecha", "Cambio", "US relacionada", "Versión"}
            if cells and cells[0] in header_cells:
                i += 1
                continue
            row_text = "  ·  ".join(c for c in cells if c)
            if row_text:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                r = p.add_run(row_text)
                set_font(r, size=9, color=C_MUTED)
            i += 1
            continue

        # Italic standalone line
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            text = stripped.strip("*").strip()
            add_body(doc, text, italic=True, color=C_MUTED)
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Generic body — strip inline markdown
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
        text = re.sub(r'\*(.+?)\*',     r'\1', text)
        text = re.sub(r'`(.+?)`',       r'\1', text)
        if text:
            add_body(doc, text)
        i += 1


# ── Main ────────────────────────────────────────────────────────────────────
def main():
    if not MD_PATH.exists():
        print(f"ERROR: No se encontró {MD_PATH}", file=sys.stderr)
        sys.exit(1)

    md_text = MD_PATH.read_text(encoding="utf-8")
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    add_cover(doc, md_text)
    parse_and_render(doc, md_text)
    doc.save(str(DOCX_OUT))
    print(f"Guardado en: {DOCX_OUT}")


if __name__ == "__main__":
    main()
